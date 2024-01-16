from docx import Document
import os

d = Document()

d.add_heading('Hello world')

p = d.add_paragraph('A paragraph')
p.add_run('\n')
p.add_run('some bold text').bold = True
p.add_run(' and some')
p.add_run(' italic text').italic = True

d.add_heading('Heading, level 1', level=1)
d.add_paragraph('Intense quote', style='Intense Quote')

items = ["item1", "item2", "item3", "item4"]

for item in items:
    d.add_paragraph("{}".format(item), style='List Bullet')
for item in items:
    d.add_paragraph("{}".format(item), style='List Number')

from docx.shared import Inches, Cm

d.add_picture('/home/guest/_repos/Scripting/Own_attempts/S2/Untitled.jpg', width=Cm(10),height=Cm(10))

d.add_page_break()

d.add_heading('Table of genes', level=2)
table = d.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'

genes = (
    ('7124', 'TNF', 'tumor necrosis factor'),
    ('4049', 'LTA', 'lymphotoxin alpha'),
    ('4050', 'LTB', 'lymphotoxin beta'),
    ('7132', 'TNFRSF1A', 'TNF receptor superfamily member 1A'),
    ('8743', 'TNFSF10', 'TNF superfamily member 10')
)

header_cells = table.rows[0].cells
header_cells[0].text = 'Gene_id'
header_cells[1].text = 'Symbol'
header_cells[2].text = 'Description'

for gene_id, symbol, desc in genes:
    row_cells = table.add_row().cells
    row_cells[0].text = gene_id
    row_cells[1].text = symbol
    row_cells[2].text = desc

# Save the document
d.save('/home/guest/_repos/Scripting/Own_attempts/S2/example.docx')

# Open the document using the default application
os.system('xdg-open /home/guest/_repos/Scripting/Own_attempts/S2/example.docx')
